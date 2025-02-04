import os  # Module pour interagir avec le système d'exploitation (gestion des chemins, répertoires, etc.)
import gzip  # Module pour lire et écrire des fichiers compressés au format gzip
import requests  # Module pour effectuer des requêtes HTTP (utilisé pour télécharger des fichiers)
import logging  # Module pour enregistrer des messages de log (suivi et débogage)
import pandas as pd  # Bibliothèque pour la manipulation et l'analyse de données sous forme de DataFrame
import matplotlib.pyplot as plt  # Bibliothèque pour créer des graphiques et visualisations
import seaborn as sns  # Bibliothèque de visualisation basée sur matplotlib (pour des graphiques statistiques)
from docx import Document  # Module pour créer et manipuler des documents Word
from docx.shared import Inches  # Permet de définir des dimensions en pouces dans un document Word
from pycaret.regression import setup, compare_models, tune_model, pull  # Fonctions de PyCaret pour la modélisation en régression

# Configuration de la journalisation pour afficher les messages d'information avec un format spécifique
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Définition du chemin de base pour stocker les fichiers téléchargés et le rapport généré
base_path = r"C:\Users\sbond\Desktop\Rapport benchmark"
# Construction du chemin complet pour le fichier des ratings (notes) compressé
ratings_path = os.path.join(base_path, "title.ratings.tsv.gz")
# Construction du chemin complet pour le fichier des informations sur les titres compressé
titles_path = os.path.join(base_path, "title.basics.tsv.gz")

# Définition des URLs des fichiers à télécharger depuis le site d'IMDB
ratings_url = "https://datasets.imdbws.com/title.ratings.tsv.gz"
titles_url = "https://datasets.imdbws.com/title.basics.tsv.gz"

try:
    # Création du dossier de base s'il n'existe pas déjà
    logging.info("Création du dossier de base...")
    os.makedirs(base_path, exist_ok=True)

    # Téléchargement des fichiers (ratings et titres) depuis les URLs
    for url, path in [(ratings_url, ratings_path), (titles_url, titles_path)]:
        logging.info(f"Téléchargement du fichier : {url}")  # Indique le début du téléchargement du fichier
        response = requests.get(url, stream=True)  # Effectue une requête GET pour télécharger le fichier
        if response.status_code == 200:  # Vérifie que le téléchargement s'est déroulé correctement (code HTTP 200)
            with open(path, "wb") as file:  # Ouvre le fichier en mode binaire pour écrire son contenu
                for chunk in response.iter_content(chunk_size=1024):  # Lit le fichier en morceaux de 1024 octets
                    file.write(chunk)  # Écrit chaque morceau dans le fichier local
            logging.info(f"Téléchargement réussi et sauvegardé dans : {path}")  # Confirme la sauvegarde du fichier
        else:
            # En cas d'erreur, lève une exception indiquant l'échec du téléchargement
            raise ConnectionError(f"Échec du téléchargement de {url}. Statut : {response.status_code}")

    # Lecture du fichier des ratings compressé en mode texte, en limitant à 1 000 000 de lignes
    logging.info("Lecture des données compressées (ratings)...")
    with gzip.open(ratings_path, 'rt', encoding='utf-8') as f:
        ratings = pd.read_csv(f, sep="\t", nrows=1000000)  # Lit le fichier TSV en DataFrame

    # Lecture du fichier des titres compressé en mode texte, en limitant à 1 000 000 de lignes
    logging.info("Lecture des données compressées (titles)...")
    with gzip.open(titles_path, 'rt', encoding='utf-8') as f:
        titles = pd.read_csv(f, sep="\t", nrows=1000000)  # Lit le fichier TSV en DataFrame

    # Fusion des deux DataFrames (ratings et titres) sur la colonne 'tconst'
    logging.info("Fusion des datasets...")
    merged_data = pd.merge(ratings, titles, on="tconst", how="inner")

    # Nettoyage des données : suppression des lignes avec des valeurs manquantes et filtrage sur le nombre de votes
    logging.info("Nettoyage des données...")
    min_votes = 50  # Seuil minimal pour le nombre de votes d'un film
    merged_data = merged_data.dropna()  # Supprime toutes les lignes contenant au moins une valeur manquante
    merged_data = merged_data[merged_data['numVotes'] > min_votes]  # Conserve uniquement les films avec plus de 50 votes
    # Sélectionne uniquement les colonnes pertinentes pour l'analyse
    merged_data = merged_data[['tconst', 'averageRating', 'numVotes', 'titleType', 'runtimeMinutes', 'startYear', 'genres']]

    # Suppression des valeurs cibles rares : conserve les ratings apparaissant au moins 2 fois
    rating_counts = merged_data['averageRating'].value_counts()  # Compte le nombre d'occurrences de chaque rating
    valid_ratings = rating_counts[rating_counts >= 2].index  # Identifie les ratings valides (au moins 2 occurrences)
    merged_data = merged_data[merged_data['averageRating'].isin(valid_ratings)]  # Filtre les lignes selon ces ratings

    # Encodage de la colonne 'genres' : création de colonnes binaires pour chaque genre séparé par une virgule
    logging.info("Encodage des genres...")
    genres_dummies = merged_data['genres'].str.get_dummies(sep=',')  # Transforme la colonne 'genres' en variables indicatrices
    merged_data = pd.concat([merged_data, genres_dummies], axis=1)  # Ajoute ces colonnes au DataFrame original
    merged_data.drop(columns=['genres'], inplace=True)  # Supprime la colonne originale 'genres'

    # Suppression de la colonne 'tconst' qui n'est plus utile après la fusion
    merged_data.drop(columns=['tconst'], inplace=True)

    # Encodage de la colonne 'titleType' en variables indicatrices (dummies)
    logging.info("Encodage de la colonne titleType...")
    titleType_dummies = merged_data['titleType'].str.get_dummies()  # Crée des colonnes binaires pour chaque type de titre
    merged_data = pd.concat([merged_data, titleType_dummies], axis=1)  # Ajoute ces colonnes au DataFrame
    merged_data.drop(columns=['titleType'], inplace=True)  # Supprime la colonne originale 'titleType'

    # Sauvegarde d'un extrait des données transformées (les 10 premières lignes) dans un fichier CSV
    sample_data = merged_data.head(10)
    sample_path = os.path.join(base_path, "extrait_donnees_transformees.csv")
    sample_data.to_csv(sample_path, index=False)  # Enregistre sans inclure l'index

    # Affichage dans le terminal de l'extrait des données pour vérification
    print("Extrait des données transformées (10 premières lignes) :")
    print(sample_data)

    # Création d'un histogramme pour visualiser la distribution de la variable cible 'averageRating'
    logging.info("Génération de l'histogramme de la variable cible...")
    plt.figure(figsize=(8, 5))  # Définit la taille de la figure
    sns.histplot(merged_data['averageRating'], kde=True, bins=20, color="skyblue")  # Trace l'histogramme avec une courbe KDE
    plt.title("Distribution de la note moyenne (averageRating)")  # Ajoute un titre au graphique
    hist_path = os.path.join(base_path, "averageRating_distribution.png")  # Définit le chemin de sauvegarde du graphique
    plt.savefig(hist_path, bbox_inches='tight')  # Sauvegarde le graphique dans un fichier PNG
    plt.close()  # Ferme la figure pour libérer les ressources

    # Création d'une carte de chaleur pour visualiser la matrice de corrélation entre les variables numériques
    logging.info("Génération de la carte de chaleur de corrélation...")
    plt.figure(figsize=(12, 10))  # Définit la taille de la figure
    numeric_cols = merged_data.select_dtypes(include=['int64', 'float64']).columns  # Sélectionne les colonnes numériques
    corr = merged_data[numeric_cols].corr()  # Calcule la matrice de corrélation
    sns.heatmap(corr, annot=True, fmt=".2f", cmap="coolwarm")  # Trace la heatmap avec annotations
    plt.title("Carte de chaleur des corrélations")  # Ajoute un titre à la heatmap
    heatmap_path = os.path.join(base_path, "correlation_heatmap.png")  # Définit le chemin de sauvegarde
    plt.savefig(heatmap_path, bbox_inches='tight')  # Sauvegarde la heatmap dans un fichier PNG
    plt.close()  # Ferme la figure

    # Initialisation de PyCaret pour la modélisation en régression avec les données nettoyées
    logging.info("Initialisation du setup PyCaret...")
    regression_setup = setup(
        data=merged_data,            # Données à utiliser pour la modélisation
        target='averageRating',      # Variable cible à prédire
        verbose=False,               # Réduit les sorties verboses de PyCaret
        session_id=123,              # Seed pour la reproductibilité
        train_size=0.8,              # Pourcentage de données utilisées pour l'entraînement
        fold=5                       # Nombre de plis pour la validation croisée
    )

    # Comparaison de plusieurs modèles de régression en utilisant PyCaret
    selected_models = ["catboost", "lightgbm", "rf", "gbr", "et"]  # Liste des modèles à considérer
    logging.info("Comparaison des modèles...")
    best_model = compare_models(include=selected_models)  # Compare et sélectionne le meilleur modèle
    best_model_name = best_model.__class__.__name__  # Extrait le nom du meilleur modèle sélectionné
    logging.info(f"Modèle sélectionné : {best_model_name}")

    # Optimisation (tuning) du modèle sélectionné pour améliorer ses performances
    tuned_model_reg = tune_model(best_model)

    # Récupération des résultats d'évaluation du modèle sous forme de DataFrame
    eval_results = pull()  # Récupère les résultats de la dernière commande PyCaret
    eval_results.index = eval_results.index.astype(str)  # Convertit les index en chaînes de caractères
    # Trie les résultats par coefficient de détermination (R²) de manière décroissante
    eval_results_sorted = eval_results.sort_values(by='R2', ascending=False).reset_index()
    # Ajoute une colonne "Rang" indiquant le classement de chaque modèle
    eval_results_sorted.insert(0, 'Rang', range(1, len(eval_results_sorted) + 1))
    # Renomme la colonne contenant le nom du modèle
    eval_results_sorted.rename(columns={eval_results_sorted.columns[1]: 'Modèle'}, inplace=True)
    eval_table_path = os.path.join(base_path, "evaluation_models.csv")  # Définit le chemin de sauvegarde du tableau d'évaluation
    eval_results_sorted.to_csv(eval_table_path, index=False)  # Enregistre le tableau en format CSV

    # Initialisation de la variable pour le chemin du graphique d'importance des variables (s'il est généré)
    importance_plot_path = None
    try:
        # Vérifie si le modèle optimisé possède l'attribut 'feature_importances_'
        if hasattr(tuned_model_reg, 'feature_importances_'):
            logging.info("Extraction de l'importance des variables...")
            feature_importance = tuned_model_reg.feature_importances_  # Récupère l'importance des variables
            # Récupère la liste des features (colonnes) en excluant la variable cible
            features = merged_data.drop(columns=['averageRating']).columns
            # Crée un DataFrame associant chaque feature à son importance
            importance_df = pd.DataFrame({'Feature': features, 'Importance': feature_importance})
            # Trie le DataFrame par ordre décroissant d'importance
            importance_df.sort_values(by='Importance', ascending=False, inplace=True)
            plt.figure(figsize=(10, 6))  # Définit la taille de la figure
            # Trace un graphique en barres montrant l'importance de chaque variable
            sns.barplot(x='Importance', y='Feature', data=importance_df, palette="viridis")
            plt.title("Importance des Variables")  # Ajoute un titre au graphique
            importance_plot_path = os.path.join(base_path, "feature_importance.png")  # Définit le chemin de sauvegarde du graphique
            plt.savefig(importance_plot_path, bbox_inches='tight')  # Sauvegarde le graphique dans un fichier PNG
            plt.close()  # Ferme la figure
        else:
            logging.info("Le modèle ne fournit pas d'attribut 'feature_importances_'.")
    except Exception as e:
        # En cas d'erreur lors de l'extraction de l'importance, log l'exception
        logging.warning("Impossible d'extraire l'importance des variables : " + str(e))

    # Création du document Word pour générer le rapport final
    logging.info("Création du rapport Word...")
    doc = Document()  # Initialise un nouveau document Word
    doc.add_heading("Benchmark des Algorithmes de Recommandation et Prédiction de Popularité", level=1)  # Ajoute le titre principal
    doc.add_paragraph("Analyse comparative des modèles de prédiction de la popularité des films en utilisant PyCaret.")  # Ajoute une description générale

    # SECTION 1 : Analyse Exploratoire

    # Sous-section pour l'analyse de la distribution de la note moyenne
    doc.add_heading("1. Analyse Exploratoire", level=2)
    doc.add_heading("1.1 Distribution de la Note Moyenne", level=3)
    doc.add_paragraph("L'histogramme ci-dessous montre la distribution de la variable cible 'averageRating'.")
    doc.add_paragraph(
        "Analyse : Ce graphique permet d'observer la répartition des notes attribuées aux films. "
        "Une distribution avec une concentration autour d'une valeur moyenne indique une homogénéité dans les évaluations. "
        "Toute dispersion ou asymétrie pourrait suggérer des biais ou des groupes distincts dans les préférences du public."
    )
    doc.add_picture(hist_path, width=Inches(5))  # Insère l'image de l'histogramme dans le rapport

    # Sous-section pour l'analyse de la carte de chaleur des corrélations
    doc.add_heading("1.2 Carte de Chaleur des Corrélations", level=3)
    doc.add_paragraph("La carte de chaleur suivante illustre les corrélations entre les variables numériques du dataset.")
    doc.add_paragraph(
        "Analyse : Cette visualisation permet d'identifier les relations linéaires entre les variables. "
        "Les zones aux coefficients élevés indiquent une forte corrélation, ce qui peut être utile pour repérer des redondances ou déterminer l'influence de certaines variables sur la cible."
    )
    doc.add_picture(heatmap_path, width=Inches(5))  # Insère l'image de la heatmap dans le rapport

    # SECTION 2 : Modélisation et Comparaison des Modèles

    doc.add_heading("2. Modélisation et Comparaison des Modèles", level=2)
    doc.add_paragraph(
        f"Le modèle sélectionné est **{best_model_name}**. Ce modèle a été choisi sur la base de ses performances, notamment en termes de R² et de RMSE, après une comparaison parmi plusieurs modèles."
    )

    # Insertion du tableau récapitulatif du classement des modèles
    doc.add_heading("2.1 Classement des Modèles", level=3)
    doc.add_paragraph("Le tableau ci-dessous présente le classement des modèles selon leurs performances :")
    eval_table = doc.add_table(rows=eval_results_sorted.shape[0] + 1, cols=eval_results_sorted.shape[1], style='Table Grid')
    # Remplissage de l'en-tête du tableau avec le nom des colonnes
    hdr_cells = eval_table.rows[0].cells
    for i, col in enumerate(eval_results_sorted.columns):
        hdr_cells[i].text = str(col)
    # Remplissage des cellules avec les données de chaque modèle et de ses métriques
    for i, row in enumerate(eval_results_sorted.itertuples(index=False), start=1):
        row_cells = eval_table.rows[i].cells
        for j, value in enumerate(row):
            row_cells[j].text = str(value)

    # SECTION 3 : Importance des Variables (affichée uniquement si disponible)
    doc.add_heading("3. Importance des Variables", level=2)
    if importance_plot_path is not None:
        doc.add_paragraph(
            "Le graphique ci-dessous présente l'importance des variables tel qu'extraites du modèle optimisé. "
            "Les variables classées en tête indiquent leur impact majeur sur la prédiction de la note moyenne."
        )
        doc.add_picture(importance_plot_path, width=Inches(5))
    else:
        doc.add_paragraph("Aucune information sur l'importance des variables n'a pu être extraite du modèle.")

    # SECTION 4 : Justification du Modèle Sélectionné
    doc.add_heading("4. Justification du Modèle Sélectionné", level=2)
    doc.add_paragraph(
        f"Le modèle sélectionné est **{best_model_name}**. "
        "Ce choix est justifié par sa capacité à maximiser le coefficient de détermination (R²) tout en minimisant l'erreur quadratique moyenne (RMSE). "
        "Les résultats obtenus après tuning indiquent que ce modèle offre un bon compromis entre performance et généralisation."
    )

    # Sauvegarde du document Word dans le chemin spécifié
    report_path = os.path.join(base_path, "rapport_benchmark.docx")
    doc.save(report_path)
    logging.info(f"Rapport sauvegardé avec succès à : {report_path}")  # Indique dans le log que le rapport a été sauvegardé

except Exception as e:
    logging.error(f"Une erreur s'est produite : {e}")  # En cas d'erreur, affiche le message d'erreur dans le log
