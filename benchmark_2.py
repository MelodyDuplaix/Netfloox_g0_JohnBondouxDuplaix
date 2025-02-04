import os
import gzip
import requests
import logging
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
from pycaret.regression import setup, compare_models, tune_model, pull

# Configuration du logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Définition des chemins de stockage
base_path = r"C:\Users\sbond\Desktop\Rapport benchmark"
ratings_path = os.path.join(base_path, "title.ratings.tsv.gz")
titles_path = os.path.join(base_path, "title.basics.tsv.gz")

# URL des fichiers IMDB
ratings_url = "https://datasets.imdbws.com/title.ratings.tsv.gz"
titles_url = "https://datasets.imdbws.com/title.basics.tsv.gz"

try:
    logging.info("Création du dossier de base...")
    os.makedirs(base_path, exist_ok=True)

    # Téléchargement des fichiers
    for url, path in [(ratings_url, ratings_path), (titles_url, titles_path)]:
        logging.info(f"Téléchargement du fichier : {url}")
        response = requests.get(url, stream=True)
        if response.status_code == 200:
            with open(path, "wb") as file:
                for chunk in response.iter_content(chunk_size=1024):
                    file.write(chunk)
            logging.info(f"Téléchargement réussi et sauvegardé dans : {path}")
        else:
            raise ConnectionError(f"Échec du téléchargement de {url}. Statut : {response.status_code}")

    # Extraction et lecture des fichiers (limité à 1 000 000 de lignes)
    logging.info("Lecture des données compressées...")
    with gzip.open(ratings_path, 'rt', encoding='utf-8') as f:
        ratings = pd.read_csv(f, sep="\t", nrows=1000000)
    with gzip.open(titles_path, 'rt', encoding='utf-8') as f:
        titles = pd.read_csv(f, sep="\t", nrows=1000000)

    # Fusion des datasets
    logging.info("Fusion des datasets...")
    merged_data = pd.merge(ratings, titles, on="tconst", how="inner")

    # Nettoyage et filtrage des données
    logging.info("Nettoyage des données...")
    min_votes = 50
    merged_data = merged_data.dropna()
    merged_data = merged_data[merged_data['numVotes'] > min_votes]
    merged_data = merged_data[['tconst', 'averageRating', 'numVotes', 'titleType', 'runtimeMinutes', 'startYear', 'genres']]

    # Suppression des valeurs cibles trop rares
    rating_counts = merged_data['averageRating'].value_counts()
    valid_ratings = rating_counts[rating_counts >= 2].index
    merged_data = merged_data[merged_data['averageRating'].isin(valid_ratings)]

    # Encodage des genres en variables indicatrices
    logging.info("Encodage des genres...")
    genres_dummies = merged_data['genres'].str.get_dummies(sep=',')
    merged_data = pd.concat([merged_data, genres_dummies], axis=1)
    merged_data.drop(columns=['genres'], inplace=True)

    # Suppression de la colonne 'tconst'
    merged_data.drop(columns=['tconst'], inplace=True)

    # Encodage de la colonne 'titleType'
    logging.info("Encodage de la colonne titleType...")
    titleType_dummies = merged_data['titleType'].str.get_dummies()
    merged_data = pd.concat([merged_data, titleType_dummies], axis=1)
    merged_data.drop(columns=['titleType'], inplace=True)

    # Sauvegarde d'un extrait des données transformées (optionnel)
    sample_data = merged_data.head(10)
    sample_path = os.path.join(base_path, "extrait_donnees_transformees.csv")
    sample_data.to_csv(sample_path, index=False)
    
    # Visualisation dans le terminal : Affichage de l'extrait des données
    print("Extrait des données transformées (10 premières lignes) :")
    print(sample_data)

    # Analyse exploratoire : Histogramme de la variable cible
    logging.info("Génération de l'histogramme de la variable cible...")
    plt.figure(figsize=(8, 5))
    sns.histplot(merged_data['averageRating'], kde=True, bins=20, color="skyblue")
    plt.title("Distribution de la note moyenne (averageRating)")
    hist_path = os.path.join(base_path, "averageRating_distribution.png")
    plt.savefig(hist_path, bbox_inches='tight')
    plt.close()

    # Analyse exploratoire : Carte de chaleur de la matrice de corrélation
    logging.info("Génération de la carte de chaleur de corrélation...")
    plt.figure(figsize=(12, 10))
    numeric_cols = merged_data.select_dtypes(include=['int64', 'float64']).columns
    corr = merged_data[numeric_cols].corr()
    sns.heatmap(corr, annot=True, fmt=".2f", cmap="coolwarm")
    plt.title("Carte de chaleur des corrélations")
    heatmap_path = os.path.join(base_path, "correlation_heatmap.png")
    plt.savefig(heatmap_path, bbox_inches='tight')
    plt.close()

    # Modélisation avec PyCaret
    logging.info("Initialisation du setup PyCaret...")
    regression_setup = setup(
        data=merged_data,
        target='averageRating',
        verbose=False,
        session_id=123,
        train_size=0.8,
        fold=5
    )

    # Sélection des modèles spécifiques
    selected_models = ["catboost", "lightgbm", "rf", "gbr", "et"]
    logging.info("Comparaison des modèles...")
    best_model = compare_models(include=selected_models)
    # Extraction propre du nom du modèle
    best_model_name = best_model.__class__.__name__
    logging.info(f"Modèle sélectionné : {best_model_name}")

    # Tuning du modèle sélectionné
    tuned_model_reg = tune_model(best_model)

    # Récupération des résultats d'évaluation
    eval_results = pull()
    eval_results.index = eval_results.index.astype(str)
    eval_results_sorted = eval_results.sort_values(by='R2', ascending=False).reset_index()
    eval_results_sorted.insert(0, 'Rang', range(1, len(eval_results_sorted) + 1))
    eval_results_sorted.rename(columns={eval_results_sorted.columns[1]: 'Modèle'}, inplace=True)
    eval_table_path = os.path.join(base_path, "evaluation_models.csv")
    eval_results_sorted.to_csv(eval_table_path, index=False)

    # Tentative d'extraction et de visualisation de l'importance des variables
    importance_plot_path = None
    try:
        # Certains modèles (basés sur les arbres) fournissent une importance des variables
        if hasattr(tuned_model_reg, 'feature_importances_'):
            logging.info("Extraction de l'importance des variables...")
            feature_importance = tuned_model_reg.feature_importances_
            # On suppose que l'ordre des colonnes correspond à merged_data (hors target)
            features = merged_data.drop(columns=['averageRating']).columns
            importance_df = pd.DataFrame({'Feature': features, 'Importance': feature_importance})
            importance_df.sort_values(by='Importance', ascending=False, inplace=True)
            plt.figure(figsize=(10, 6))
            sns.barplot(x='Importance', y='Feature', data=importance_df, palette="viridis")
            plt.title("Importance des Variables")
            importance_plot_path = os.path.join(base_path, "feature_importance.png")
            plt.savefig(importance_plot_path, bbox_inches='tight')
            plt.close()
        else:
            logging.info("Le modèle ne fournit pas d'attribut 'feature_importances_'.")
    except Exception as e:
        logging.warning("Impossible d'extraire l'importance des variables : " + str(e))

    # Création du rapport Word avec analyses graphiques
    logging.info("Création du rapport Word...")
    doc = Document()
    doc.add_heading("Benchmark des Algorithmes de Recommandation et Prédiction de Popularité", level=1)
    doc.add_paragraph("Analyse comparative des modèles de prédiction de la popularité des films en utilisant PyCaret.")

    # Section 1 : Analyse Exploratoire
    doc.add_heading("1. Analyse Exploratoire", level=2)

    # Histogramme de la note moyenne
    doc.add_heading("1.1 Distribution de la Note Moyenne", level=3)
    doc.add_paragraph("L'histogramme ci-dessous montre la distribution de la variable cible 'averageRating'.")
    doc.add_paragraph(
        "Analyse : Ce graphique permet d'observer la répartition des notes attribuées aux films. "
        "Une distribution avec une concentration autour d'une valeur moyenne indique une homogénéité dans les évaluations. "
        "Toute dispersion ou asymétrie pourrait suggérer des biais ou des groupes distincts dans les préférences du public."
    )
    doc.add_picture(hist_path, width=Inches(5))

    # Carte de chaleur des corrélations
    doc.add_heading("1.2 Carte de Chaleur des Corrélations", level=3)
    doc.add_paragraph("La carte de chaleur suivante illustre les corrélations entre les variables numériques du dataset.")
    doc.add_paragraph(
        "Analyse : Cette visualisation permet d'identifier les relations linéaires entre les variables. "
        "Les zones aux coefficients élevés indiquent une forte corrélation, ce qui peut être utile pour repérer des redondances ou déterminer l'influence de certaines variables sur la cible."
    )
    doc.add_picture(heatmap_path, width=Inches(5))
    
    # Section 2 : Modélisation et Comparaison des Modèles
    doc.add_heading("2. Modélisation et Comparaison des Modèles", level=2)
    doc.add_paragraph(
        f"Le modèle sélectionné est **{best_model_name}**. Ce modèle a été choisi sur la base de ses performances, notamment en termes de R² et de RMSE, après une comparaison parmi plusieurs modèles."
    )
    
    # Insertion du tableau de comparaison des modèles
    doc.add_heading("2.1 Classement des Modèles", level=3)
    doc.add_paragraph("Le tableau ci-dessous présente le classement des modèles selon leurs performances :")
    eval_table = doc.add_table(rows=eval_results_sorted.shape[0] + 1, cols=eval_results_sorted.shape[1], style='Table Grid')
    # Remplissage de l'en-tête
    hdr_cells = eval_table.rows[0].cells
    for i, col in enumerate(eval_results_sorted.columns):
        hdr_cells[i].text = str(col)
    # Remplissage des lignes du tableau
    for i, row in enumerate(eval_results_sorted.itertuples(index=False), start=1):
        row_cells = eval_table.rows[i].cells
        for j, value in enumerate(row):
            row_cells[j].text = str(value)
    
    # Section 3 : Importance des Variables (si disponible)
    doc.add_heading("3. Importance des Variables", level=2)
    if importance_plot_path is not None:
        doc.add_paragraph(
            "Le graphique ci-dessous présente l'importance des variables tel qu'extraites du modèle optimisé. "
            "Les variables classées en tête indiquent leur impact majeur sur la prédiction de la note moyenne."
        )
        doc.add_picture(importance_plot_path, width=Inches(5))
    else:
        doc.add_paragraph("Aucune information sur l'importance des variables n'a pu être extraite du modèle.")
    
    # Section 4 : Justification du Modèle Sélectionné
    doc.add_heading("4. Justification du Modèle Sélectionné", level=2)
    doc.add_paragraph(
        f"Le modèle sélectionné est **{best_model_name}**. "
        "Ce choix est justifié par sa capacité à maximiser le coefficient de détermination (R²) tout en minimisant l'erreur quadratique moyenne (RMSE). "
        "Les résultats obtenus après tuning indiquent que ce modèle offre un bon compromis entre performance et généralisation."
    )
    
    # Sauvegarde du rapport Word
    report_path = os.path.join(base_path, "rapport_benchmark.docx")
    doc.save(report_path)
    logging.info(f"Rapport sauvegardé avec succès à : {report_path}")

except Exception as e:
    logging.error(f"Une erreur s'est produite : {e}")
